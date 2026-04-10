import json
from pathlib import Path

from docx import Document

from scripts.pipeline import PipelineConfig, convert_markdown_to_docx


class FakeNormalizer:
    def to_latex(self, expression: str) -> str:
        if "E =" in expression:
            return r"E = (1 - \frac{RH}{100}) \times (1 + 0.35v)"
        if "a/b" in expression:
            return r"\frac{a}{b}"
        return expression


def _write_test_xsl(path: Path) -> None:
    xsl = """<?xml version="1.0" encoding="UTF-8"?>
<xsl:stylesheet version="1.0"
    xmlns:xsl="http://www.w3.org/1999/XSL/Transform"
    xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
  <xsl:output method="xml" omit-xml-declaration="yes"/>
  <xsl:template match="/">
    <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r><m:t><xsl:value-of select="string(.)"/></m:t></m:r>
    </m:oMath>
  </xsl:template>
</xsl:stylesheet>
"""
    path.write_text(xsl, encoding="utf-8")


def test_pipeline_generates_docx_with_formula_and_report(tmp_path: Path):
    md_path = tmp_path / "input.md"
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "output.docx"
    log_path = tmp_path / "report.json"
    xsl_path = tmp_path / "mml2omml.xsl"
    _write_test_xsl(xsl_path)

    md_path.write_text(
        "# 标题\n\n正文 `E = (1 - RH / 100) × (1 + 0.35v)` 和 `print(x)`。\n\n- 列表项 `a/b`\n",
        encoding="utf-8",
    )

    template_doc = Document()
    template_doc.styles["Normal"].font.name = "Times New Roman"
    template_doc.save(str(template_path))

    config = PipelineConfig(mml2omml_xsl=str(xsl_path))
    result = convert_markdown_to_docx(
        input_md_path=str(md_path),
        template_docx_path=str(template_path),
        output_docx_path=str(output_path),
        config=config,
        log_path=str(log_path),
        normalizer=FakeNormalizer(),
    )

    assert output_path.exists()
    assert log_path.exists()
    assert result["formula_success"] >= 1
    assert result["formula_failed"] == 0

    report_from_file = json.loads(log_path.read_text(encoding="utf-8"))
    assert report_from_file["formula_success"] >= 1

    out_doc = Document(str(output_path))
    assert out_doc.paragraphs[0].style.name in {"Heading 1", "标题 1"}
    assert any("`print(x)`" in p.text for p in out_doc.paragraphs)
    assert "m:oMath" in out_doc._element.xml
